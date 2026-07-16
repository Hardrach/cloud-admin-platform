from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
from datetime import datetime
import os
import subprocess
import textwrap


ROOT = Path(r"D:\Projet Stage Deweb\New projet\Projects\cloud-admin-platform")
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "assets"
OUT_DOCX = OUT_DIR / "Cloud_Admin_Platform_Rapport_Technique_Complet.docx"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

SCREENSHOTS = [
    (r"C:\Users\yassi\AppData\Local\Temp\codex-clipboard-bc12bf20-4ce3-4636-a19d-caf44298824c.png", "Vue reseau en light mode avec interfaces et diagnostics."),
    (r"C:\Users\yassi\AppData\Local\Temp\codex-clipboard-adb278cc-7de3-40ef-87f6-f79cc966b848.png", "Inventaire Docker avec conteneurs reels et actions d'administration."),
    (r"C:\Users\yassi\AppData\Local\Temp\codex-clipboard-b3286f9c-b618-4448-a701-41a648a25450.png", "Page Firewall et regles ACL issues de l'API."),
    (r"C:\Users\yassi\AppData\Local\Temp\codex-clipboard-4be62aea-6190-4f5f-bd6d-0232ceb35159.png", "Page Settings apres optimisation de la colonne droite."),
    (r"C:\Users\yassi\AppData\Local\Temp\codex-clipboard-9b096dfa-b5e7-473b-9d89-6f46707771f5.png", "Navbar finale sans doublon de branding."),
]


def run(cmd):
    try:
        return subprocess.check_output(cmd, cwd=ROOT, text=True, stderr=subprocess.STDOUT).strip()
    except Exception:
        return ""


ENDPOINTS = [
    ("GET", "/api/dashboard", "Synthese infrastructure, systeme, Docker, Git et Terraform", "psutil, Docker SDK, Git CLI, Terraform CLI"),
    ("GET", "/api/vms", "Inventaire et etat de la machine virtuelle Azure", "Azure CLI, fallback local"),
    ("POST", "/api/vms/{name}/start", "Demarrage VM", "Azure CLI"),
    ("POST", "/api/vms/{name}/stop", "Arret VM", "Azure CLI"),
    ("POST", "/api/vms/{name}/restart", "Redemarrage VM", "Azure CLI"),
    ("POST", "/api/vms/{name}/deallocate", "Deallocation VM", "Azure CLI"),
    ("GET", "/api/docker", "Inventaire conteneurs Docker", "Docker SDK, docker CLI"),
    ("GET", "/api/docker/stats", "CPU/RAM par conteneur", "Docker stats"),
    ("POST", "/api/docker/{container}/start", "Demarrage conteneur", "Docker SDK"),
    ("POST", "/api/docker/{container}/stop", "Arret conteneur", "Docker SDK"),
    ("POST", "/api/docker/{container}/restart", "Redemarrage conteneur", "Docker SDK"),
    ("GET", "/api/docker/{container}/logs", "Logs conteneur", "Docker SDK"),
    ("GET", "/api/networks", "Interfaces, IP, DNS, gateway, RX/TX", "psutil, Docker SDK"),
    ("GET", "/api/storage", "Volumes Docker et stockage hote", "Docker SDK, shutil.disk_usage"),
    ("GET", "/api/metrics", "CPU, RAM, load, swap, disk IO, reseau", "psutil"),
    ("GET", "/api/logs", "Logs systeme et conteneurs", "Docker logs, fichiers systeme"),
    ("GET", "/api/alerts", "Alertes issues des seuils systeme", "psutil, Docker, Git"),
    ("GET", "/api/firewall", "Statut UFW et regles ACL", "ufw CLI"),
    ("GET", "/api/ssh-keys", "Cles SSH autorisees", "~/.ssh"),
    ("GET", "/api/iam", "Utilisateurs Linux", "pwd, /etc/passwd"),
    ("GET", "/api/terraform", "Version, fichiers, state, workspace", "Terraform CLI"),
    ("POST", "/api/terraform/plan", "Execution terraform plan", "Terraform CLI"),
    ("POST", "/api/terraform/apply", "Execution terraform apply", "Terraform CLI"),
    ("POST", "/api/terraform/destroy", "Execution terraform destroy", "Terraform CLI"),
    ("GET", "/api/docker-compose", "Stack Docker Compose", "Docker Compose CLI"),
    ("GET", "/api/github", "Etat Git local et historique commits", "Git CLI"),
    ("POST", "/api/github/fetch", "git fetch origin", "Git CLI"),
    ("POST", "/api/github/pull", "git pull origin branch", "Git CLI"),
    ("POST", "/api/github/push", "git push origin branch", "Git CLI"),
    ("POST", "/api/github/commit", "Commit workspace", "Git CLI"),
]

TECH = [
    ("Frontend", "React", "Interface SPA, composants, etat utilisateur"),
    ("Frontend", "Framer Motion", "Transitions de pages et micro-interactions"),
    ("Frontend", "Recharts", "Graphiques de monitoring"),
    ("Frontend", "Axios", "Client REST vers FastAPI"),
    ("Backend", "FastAPI", "API REST modulaire"),
    ("Backend", "psutil", "Metriques CPU, RAM, disque, reseau"),
    ("Backend", "Docker SDK", "Inventaire et actions conteneurs"),
    ("Backend", "subprocess", "Pont vers Git, Terraform, Azure CLI, UFW"),
    ("DevOps", "Docker Compose", "Orchestration backend + PostgreSQL"),
    ("Cloud", "Azure VM", "Hebergement backend et outils systeme"),
    ("Edge", "Nginx", "Reverse proxy HTTPS"),
    ("DNS/SSL", "DuckDNS + Let's Encrypt", "Nom de domaine et certificats"),
    ("Deploy", "Vercel", "Hebergement frontend React"),
]


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_text(hdr[i], h, True, "0B2545")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    for line in code.strip().splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cliquez droit puis Mettre a jour le champ pour actualiser la table des matieres."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def diagram(path, title, boxes, arrows):
    img = Image.new("RGB", (1600, 900), "#F8FAFC")
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 42)
        font = ImageFont.truetype("arial.ttf", 24)
        font_small = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        font_title = font = font_small = None
    d.text((60, 45), title, fill="#0B2545", font=font_title)
    for key, label, xy, fill in boxes:
        x, y, w, h = xy
        d.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline="#CBD5E1", width=3)
        lines = textwrap.wrap(label, 18)
        for idx, line in enumerate(lines):
            d.text((x + 28, y + 28 + idx * 30), line, fill="#0F172A", font=font)
    centers = {key: (xy[0] + xy[2] // 2, xy[1] + xy[3] // 2) for key, _, xy, _ in boxes}
    for a, b, label in arrows:
        ax, ay = centers[a]
        bx, by = centers[b]
        d.line((ax, ay, bx, by), fill="#2563EB", width=5)
        d.ellipse((bx - 8, by - 8, bx + 8, by + 8), fill="#2563EB")
        if label:
            d.text(((ax + bx) // 2 + 8, (ay + by) // 2 - 24), label, fill="#334155", font=font_small)
    img.save(path)


def make_diagrams():
    architecture = ASSET_DIR / "architecture.png"
    diagram(
        architecture,
        "Architecture globale Cloud Admin Platform",
        [
            ("user", "Utilisateur / Navigateur", (50, 220, 250, 110), "#DBEAFE"),
            ("vercel", "Frontend React sur Vercel", (380, 220, 260, 110), "#CCFBF1"),
            ("dns", "DuckDNS + HTTPS", (720, 220, 230, 110), "#FEF3C7"),
            ("nginx", "Nginx Reverse Proxy", (1030, 220, 240, 110), "#EDE9FE"),
            ("api", "FastAPI Backend Azure VM", (1320, 220, 240, 110), "#DCFCE7"),
            ("docker", "Docker / Compose", (560, 530, 250, 110), "#E0F2FE"),
            ("linux", "Linux / psutil / UFW / SSH", (880, 530, 280, 110), "#FCE7F3"),
            ("devops", "Git / Terraform / Azure CLI", (1200, 530, 280, 110), "#F1F5F9"),
        ],
        [
            ("user", "vercel", "SPA"),
            ("vercel", "dns", "REST HTTPS"),
            ("dns", "nginx", "443"),
            ("nginx", "api", "proxy"),
            ("api", "docker", "SDK"),
            ("api", "linux", "system"),
            ("api", "devops", "CLI"),
        ],
    )
    dataflow = ASSET_DIR / "dataflow.png"
    diagram(
        dataflow,
        "Flux de donnees d'une page React vers le systeme",
        [
            ("page", "Page React", (70, 240, 210, 100), "#DBEAFE"),
            ("axios", "Axios Service", (360, 240, 210, 100), "#CCFBF1"),
            ("route", "Endpoint FastAPI", (650, 240, 230, 100), "#DCFCE7"),
            ("collector", "Collecteur Python", (960, 240, 250, 100), "#FEF3C7"),
            ("source", "Source reelle: OS / Docker / Git", (1300, 240, 250, 100), "#EDE9FE"),
            ("json", "JSON normalise", (650, 520, 230, 100), "#F1F5F9"),
            ("ui", "Composants UI", (360, 520, 210, 100), "#FCE7F3"),
        ],
        [
            ("page", "axios", "call"),
            ("axios", "route", "GET/POST"),
            ("route", "collector", "Python"),
            ("collector", "source", "read/action"),
            ("collector", "json", "response"),
            ("json", "ui", "render"),
            ("ui", "page", "state"),
        ],
    )
    deploy = ASSET_DIR / "deployment.png"
    diagram(
        deploy,
        "Workflow de deploiement et exploitation",
        [
            ("local", "Developpement local", (90, 220, 250, 110), "#DBEAFE"),
            ("git", "GitHub main", (420, 220, 230, 110), "#F1F5F9"),
            ("vercel", "Vercel build frontend", (730, 220, 260, 110), "#CCFBF1"),
            ("azure", "Azure VM pull/build", (1060, 220, 260, 110), "#EDE9FE"),
            ("compose", "docker compose up -d --build", (1060, 520, 310, 110), "#DCFCE7"),
            ("nginx", "Nginx + Certbot", (710, 520, 260, 110), "#FEF3C7"),
            ("prod", "Production HTTPS", (380, 520, 250, 110), "#FCE7F3"),
        ],
        [
            ("local", "git", "commit/push"),
            ("git", "vercel", "auto deploy"),
            ("git", "azure", "pull"),
            ("azure", "compose", "build"),
            ("compose", "nginx", "serve API"),
            ("nginx", "prod", "443"),
            ("vercel", "prod", "frontend"),
        ],
    )
    return [architecture, dataflow, deploy]


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11.5, "1F4D78")]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.bold = True
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)
    footer = section.footer.paragraphs[0]
    footer.text = "Cloud Admin Platform - Rapport technique complet"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Cloud Admin Platform")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor(11, 37, 69)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rapport Technique Complet")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(46, 116, 181)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Projet de Fin d'Etudes - Portfolio DevOps / Cloud Administration")
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = RGBColor(71, 85, 105)
    add_table(doc, ["Element", "Valeur"], [
        ("Projet", "Plateforme Web Full Stack de supervision et administration Cloud"),
        ("Architecture", "React + FastAPI + Docker + Azure VM + Nginx + DuckDNS + Vercel"),
        ("Version du rapport", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Repository", "https://github.com/Hardrach/cloud-admin-platform.git"),
    ], [1.7, 4.8])
    doc.add_page_break()


def h1(doc, title):
    doc.add_heading(title, level=1)


def h2(doc, title):
    doc.add_heading(title, level=2)


def h3(doc, title):
    doc.add_heading(title, level=3)


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def section_page(doc, title, content, bullets_list=None, table=None, code=None):
    h1(doc, title)
    for block in content:
        para(doc, block)
    if bullets_list:
        bullets(doc, bullets_list)
    if table:
        add_table(doc, table[0], table[1], table[2] if len(table) > 2 else None)
    if code:
        add_code(doc, code)
    doc.add_page_break()


def build():
    diagrams = make_diagrams()
    doc = Document()
    style_doc(doc)
    add_title_page(doc)

    h1(doc, "Table des matieres")
    add_toc(doc)
    doc.add_page_break()

    git_branch = run(["git", "branch", "--show-current"]) or "main"
    git_head = run(["git", "rev-parse", "--short", "HEAD"]) or "N/A"
    file_count = len(list((ROOT / "frontend" / "src").rglob("*.*"))) + len(list((ROOT / "backend").rglob("*.*")))

    section_page(doc, "1. Presentation generale du projet", [
        "Cloud Admin Platform est une plateforme Web Full Stack permettant de superviser et administrer une infrastructure Cloud depuis une interface moderne.",
        "Le projet centralise les informations de machines virtuelles, conteneurs Docker, reseau, stockage, logs, alertes, firewall, SSH, IAM, Terraform et GitHub.",
        "L'objectif PFE est de demontrer une competence complete en developpement frontend, backend, DevOps, cloud, securite et exploitation systeme."
    ], ["Nom: Cloud Admin Platform", "Type: plateforme SaaS DevOps / Cloud Administration", "Branche analysee: " + git_branch, "Commit analyse: " + git_head, f"Fichiers applicatifs analyses: {file_count}"])

    section_page(doc, "2. Objectifs fonctionnels atteints", [
        "Le projet vise une supervision temps reel avec un frontend lisible, des donnees collectees depuis le systeme et des actions d'administration exposees par API.",
        "Les objectifs couvrent l'observabilite, l'administration Docker, la lecture de l'etat Git, l'inspection Terraform, la securite Linux et le monitoring systeme."
    ], [
        "Superviser la machine Azure et son etat.",
        "Superviser Docker, Docker Compose et les conteneurs.",
        "Lire les logs, alertes et metriques CPU/RAM/disque/reseau.",
        "Scanner les utilisateurs Linux, cles SSH et regles UFW.",
        "Afficher GitHub, Terraform, Storage, Networks et Dashboard depuis API.",
        "Fournir Dark Mode, Light Mode et mode automatique persistant."
    ])

    h1(doc, "3. Architecture generale")
    para(doc, "L'architecture separe clairement le frontend React, le backend FastAPI, les outils systeme et les services d'infrastructure.")
    doc.add_picture(str(diagrams[0]), width=Inches(6.4))
    add_caption(doc, "Figure 1 - Architecture globale frontend, backend, Azure, Docker, Nginx, DuckDNS et Vercel.")
    para(doc, "Cette architecture evite de mettre la logique systeme dans le navigateur. Le frontend consomme uniquement des endpoints REST securises.")
    doc.add_page_break()

    h1(doc, "4. Flux de donnees")
    para(doc, "Chaque page React appelle un service Axios, qui interroge un endpoint FastAPI. Le backend interroge ensuite les sources reelles: psutil, Docker SDK, Git, Terraform, UFW ou Azure CLI.")
    doc.add_picture(str(diagrams[1]), width=Inches(6.4))
    add_caption(doc, "Figure 2 - Flux de donnees d'une page React vers les sources systeme.")
    doc.add_page_break()

    h1(doc, "5. Workflow de deploiement")
    para(doc, "Le workflow DevOps repose sur GitHub comme source de verite. Le frontend peut etre deploye par Vercel, tandis que le backend est reconstruit sur la VM Azure avec Docker Compose.")
    doc.add_picture(str(diagrams[2]), width=Inches(6.4))
    add_caption(doc, "Figure 3 - Workflow local, GitHub, Vercel, Azure VM et Docker Compose.")
    doc.add_page_break()

    section_page(doc, "6. Technologies utilisees", [
        "La stack associe une interface React moderne et un backend Python capable d'interroger des donnees bas niveau. Les outils DevOps assurent l'orchestration, le deploiement et la securisation."
    ], table=(["Couche", "Technologie", "Role"], TECH, [1.2, 1.7, 3.6]))

    section_page(doc, "7. Structure du repository", [
        "Le projet est organise en dossiers distincts afin de separer l'application frontend, le backend API, l'infrastructure et les fichiers de deploiement."
    ], code="""
cloud-admin-platform/
  backend/
    main.py
    requirements.txt
    Dockerfile
  frontend/
    src/
      components/
      pages/
      services/api.js
    package.json
  infrastructure/terraform/
  docker-compose.yml
  README.md
""")

    section_page(doc, "8. Backend FastAPI", [
        "Le backend expose une API REST et agit comme couche d'abstraction entre l'interface Web et les outils systeme.",
        "FastAPI permet de definir rapidement des endpoints typés, testables, et faciles a consommer depuis React."
    ], [
        "Lecture systeme avec psutil.",
        "Controle Docker via Docker SDK.",
        "Execution Git, Terraform, UFW et Azure CLI via subprocess.",
        "Gestion CORS pour la communication frontend-backend.",
        "Fallbacks controles pour les environnements locaux."
    ], code="""
@app.get("/api/dashboard")
def get_dashboard() -> dict:
    return {
        "system": {
            "cpu": round(psutil.cpu_percent(), 0),
            "memory": round(psutil.virtual_memory().percent, 0)
        }
    }
""")

    section_page(doc, "9. Frontend React", [
        "Le frontend est une SPA React composee de pages metier, composants reutilisables, transitions et design system global.",
        "Les pages sont connectees a services/api.js pour garantir que les donnees viennent de l'API et non de tableaux statiques."
    ], [
        "Dashboard",
        "Virtual Machines",
        "Docker Containers",
        "Networks",
        "Storage",
        "Metrics",
        "Logs",
        "Alerts",
        "Firewall",
        "SSH Keys",
        "IAM",
        "Terraform",
        "Docker Compose",
        "GitHub",
        "Settings et Profile"
    ])

    h1(doc, "10. Captures d'ecran de l'interface")
    for idx, (path, caption) in enumerate(SCREENSHOTS, start=1):
        if Path(path).exists():
            try:
                doc.add_picture(path, width=Inches(6.3))
                add_caption(doc, f"Figure {idx + 3} - {caption}")
            except Exception:
                para(doc, f"Capture non integree: {path}")
    doc.add_page_break()

    h1(doc, "11. Tableau recapitulatif des endpoints")
    add_table(doc, ["Methode", "Endpoint", "Fonction", "Source"], ENDPOINTS[:15], [0.75, 1.8, 2.35, 1.6])
    doc.add_page_break()
    h1(doc, "12. Tableau recapitulatif des endpoints - suite")
    add_table(doc, ["Methode", "Endpoint", "Fonction", "Source"], ENDPOINTS[15:], [0.75, 1.8, 2.35, 1.6])
    doc.add_page_break()

    pages = [
        ("13. Dashboard", ["Le Dashboard centralise les indicateurs systeme et applicatifs: CPU, RAM, disque, conteneurs, base de donnees, version Docker, version Terraform, hostname, kernel, uptime et branche Git.", "Cette page donne une vision rapide de la sante de la plateforme."], ["getDashboard()", "psutil.cpu_percent()", "psutil.virtual_memory()", "Docker SDK", "Git CLI"]),
        ("14. Virtual Machines", ["La page Virtual Machines affiche l'etat de la VM Azure et expose des actions start, stop, restart et deallocate.", "Le frontend a ete corrige pour lire le format liste renvoye par /api/vms."], ["GET /api/vms", "POST /api/vms/{name}/start", "POST /api/vms/{name}/stop"]),
        ("15. Docker Containers", ["La page Docker affiche les conteneurs reels et les actions d'administration.", "Les boutons Start, Stop, Restart et Logs utilisent des endpoints FastAPI relies au Docker daemon."], ["GET /api/docker", "GET /api/docker/stats", "POST /api/docker/{container}/stop"]),
        ("16. Docker Compose", ["Cette page inspecte la stack Compose, les conteneurs rattaches et les definitions disponibles.", "Elle aide a comprendre l'orchestration locale et production."], ["docker compose version", "docker compose ps", "containers labels"]),
        ("17. Networks", ["La page reseau expose interfaces, IP, MAC, DNS, gateway, RX et TX.", "Les donnees proviennent de psutil et Docker networks."], ["psutil.net_if_stats", "psutil.net_if_addrs", "psutil.net_io_counters"]),
        ("18. Storage", ["La page Storage presente l'usage disque et les volumes Docker.", "Elle combine shutil.disk_usage et les volumes Docker."], ["shutil.disk_usage", "Docker volumes"]),
        ("19. Metrics", ["Metrics affiche l'historique CPU/RAM ainsi que les donnees swap, disk IO, reseau et load average.", "Le backend maintient un historique en memoire pour rendre les graphes lisibles."], ["cpu_history", "memory_history", "network_history"]),
        ("20. Logs", ["La page Logs agrege les logs Docker et les logs systeme lorsque disponibles.", "Elle inclut filtres de niveau et recherche textuelle."], ["container.logs", "/var/log/auth.log", "/var/log/syslog"]),
        ("21. Alerts", ["Alerts transforme les signaux systeme en alertes exploitables.", "Les alertes couvrent ressources, conteneurs, Git, Terraform et base de donnees."], ["seuil CPU", "seuil RAM", "Git dirty", "Terraform absent"]),
        ("22. Firewall", ["Firewall lit les regles UFW et les transforme en ACL lisibles.", "La page affiche port, protocole, action, source et commentaire."], ["ufw status verbose", "ufw status numbered"]),
        ("23. SSH Keys", ["SSH Keys analyse les cles publiques presentes sur l'hote.", "Elle presente utilisateur, fingerprint et date d'import."], ["~/.ssh", "*.pub"]),
        ("24. IAM", ["IAM liste les utilisateurs Linux et leurs metadonnees.", "Cette page est utile pour la verification systeme et securite."], ["pwd.getpwall", "UID", "home", "shell"]),
        ("25. Terraform", ["Terraform inspecte l'installation, les fichiers .tf, le workspace, les outputs et le state.", "Des actions plan, apply et destroy sont exposees via API."], ["terraform -version", "terraform state list", "terraform plan"]),
        ("26. GitHub", ["GitHub affiche remote, branche, historique, statut, ahead/behind et taille du repository.", "La section commit locale a ete retiree pour simplifier l'interface."], ["git status", "git log", "git remote"]),
        ("27. Settings", ["Settings gere les preferences utilisateur et la configuration API.", "Le mode theme automatique est persistant via localStorage et suit l'heure et le systeme."], ["localStorage", "prefers-color-scheme", "REACT_APP_API_URL"]),
        ("28. Profile", ["Profile represente la zone utilisateur et les donnees de session.", "Elle reste volontairement plus statique car elle attend une future authentification."], ["avatar", "email", "role"]),
    ]

    for title, paragraphs, points in pages:
        section_page(doc, title, paragraphs, points)

    section_page(doc, "29. Configuration des environnements", [
        "Le frontend utilise les variables React pour choisir l'API selon l'environnement.",
        "En developpement, l'API pointe sur localhost. En production, elle doit pointer vers le domaine HTTPS ou l'IP publique."
    ], code="""
REACT_APP_API_URL=http://localhost:8000
REACT_APP_API_URL=https://cloudadminyassine.duckdns.org
""")

    section_page(doc, "30. Securisation HTTPS", [
        "La securisation repose sur DuckDNS, Nginx, Certbot et Let's Encrypt.",
        "Nginx agit comme reverse proxy entre l'Internet public et FastAPI."
    ], ["Port 80 pour validation HTTP.", "Port 443 pour HTTPS.", "Proxy vers localhost:8000.", "Renouvellement automatique Certbot."], code="""
server {
  listen 443 ssl;
  server_name cloudadminyassine.duckdns.org;
  location / {
    proxy_pass http://localhost:8000;
  }
}
""")

    section_page(doc, "31. Docker et orchestration", [
        "Docker Compose lance PostgreSQL et le backend FastAPI. Le socket Docker est monte afin que l'API puisse inspecter et administrer les conteneurs.",
        "Cette approche est puissante mais impose de controler les permissions et l'exposition du backend."
    ], code="""
docker compose up -d --build
docker compose ps
docker compose logs -f backend
docker compose down
""")

    section_page(doc, "32. Workflow Git et GitHub", [
        "GitHub centralise le code source. Le flux consiste a developper localement, valider par build, commit, push, puis deployer.",
        "Le backend expose aussi des informations Git pour que l'interface affiche l'etat du repository."
    ], ["Developpement local", "Commit", "Push GitHub", "Pull Azure", "Docker build", "Production"])

    section_page(doc, "33. Workflow Vercel", [
        "Vercel peut reconstruire automatiquement le frontend a chaque push.",
        "La variable d'environnement API doit pointer vers le backend public HTTPS."
    ], ["Push GitHub", "Build React", "Publication Vercel", "Communication HTTPS avec Azure backend"])

    section_page(doc, "34. Design system et UI", [
        "Le design system repose sur variables CSS, composants reutilisables, light/dark mode et animations.",
        "La plateforme a evolue vers une interface SaaS plus professionnelle avec micro-interactions, transitions, fonds animes et cartes interactives."
    ], ["Variables CSS", "Theme clair/sombre/auto", "PageTransition Framer Motion", "Skeletons", "Toast", "ConfirmDialog", "DataTable"])

    section_page(doc, "35. Theme automatique et experience utilisateur", [
        "Le mode automatique applique le light mode le jour et le dark mode la nuit, tout en respectant la preference systeme.",
        "Le choix utilisateur est sauvegarde dans localStorage afin que la preference soit conservee par navigateur."
    ], ["Auto", "Light", "Dark", "Premier acces avec choix", "Persistance par utilisateur"])

    section_page(doc, "36. Problèmes rencontres", [
        "Le projet a rencontre plusieurs problemes typiques d'une plateforme DevOps: proprietes Git dans Docker, mixed content HTTPS, mapping API/frontend, permissions Docker, CORS et configuration Nginx.",
        "Ces problemes ont ete resolus progressivement par audit et refactorisation."
    ], ["fatal: detected dubious ownership", "N/A dans les pages", "Docker mock non controle", "Light mode avec fond sombre", "Endpoints et formats JSON heterogenes", "Push GitHub et rebase"])

    section_page(doc, "37. Solutions apportees", [
        "Les corrections ont porte sur la normalisation des donnees, la suppression des fausses donnees critiques, l'enrichissement des endpoints et l'amelioration UX.",
        "Le backend renvoie maintenant des donnees plus completes pour les pages reseau, Docker, GitHub et Firewall."
    ], ["git config safe.directory", "HTTPS via Nginx", "Mapping React corrige", "Suppression de fallbacks Docker trompeurs", "Mode theme persistant", "Build et push GitHub valides"])

    section_page(doc, "38. Audit des donnees dynamiques", [
        "Les pages metier sont connectees aux endpoints API. Settings et Profile restent des zones de configuration/session qui seront enrichies avec une authentification future.",
        "L'audit montre que les donnees principales proviennent de sources reelles: Linux, Docker, Git, Terraform, Azure CLI et UFW."
    ], table=(["Domaine", "Source reelle", "Page"], [
        ("Systeme", "psutil", "Dashboard, Metrics, Storage, Networks"),
        ("Docker", "Docker SDK", "Docker, Logs, Compose"),
        ("Git", "Git CLI", "GitHub, Dashboard"),
        ("Terraform", "Terraform CLI", "Terraform, Dashboard"),
        ("Firewall", "UFW CLI", "Firewall"),
        ("IAM/SSH", "OS Linux", "IAM, SSH Keys"),
    ], [1.5, 2.4, 2.6]))

    section_page(doc, "39. Securite et risques", [
        "Le projet manipule des capacites sensibles: Docker socket, commandes systeme, Terraform apply/destroy et controle VM.",
        "En production, ces actions doivent etre protegees par authentification, autorisation RBAC, journalisation et confirmation forte."
    ], ["JWT ou OAuth2", "RBAC par role", "Audit log", "Rate limiting", "Protection endpoints destructifs", "Secrets hors repository"])

    section_page(doc, "40. Tests et validation", [
        "Les validations realisees incluent build React, compilation Python, tests manuels API, actions Docker et verification GitHub.",
        "Le workflow actuel peut etre renforce par GitHub Actions et tests automatises."
    ], ["npm.cmd run build", "python -m py_compile backend/main.py", "Invoke-WebRequest /health", "Stop/Start postgres-db via API", "git push origin main"])

    section_page(doc, "41. Commandes de lancement", [
        "Les commandes suivantes permettent de lancer le projet en developpement, en Docker et en verification production."
    ], code="""
cd backend
uvicorn main:app --reload

cd frontend
npm install
npm start

docker compose up -d --build
docker compose logs -f backend
""")

    section_page(doc, "42. Competences demontrees", [
        "Ce projet demontre une polyvalence complete en developpement Full Stack, exploitation Linux, cloud Azure, DevOps et design d'interface.",
        "Il constitue une base solide pour un PFE, un portfolio GitHub ou un entretien technique."
    ], ["React", "FastAPI", "Python", "Docker", "Docker Compose", "Linux", "Azure", "Git/GitHub", "Terraform", "Nginx", "Let's Encrypt", "DuckDNS", "Vercel", "REST API", "Monitoring"])

    section_page(doc, "43. Pistes d'amelioration", [
        "La plateforme peut evoluer vers un vrai produit SaaS d'administration Cloud avec authentification, temps reel et monitoring avance.",
        "Les evolutions prioritaires concernent la securite, l'observabilite, l'automatisation CI/CD et l'UX avancee."
    ], ["Authentification JWT", "Gestion des roles", "WebSockets temps reel", "Prometheus/Grafana", "Kubernetes", "GitHub Actions", "Audit logs", "Notifications temps reel", "Three.js / GSAP"])

    section_page(doc, "44. Roadmap proposee", [
        "La roadmap ci-dessous structure le passage d'un PFE fonctionnel vers une plateforme plus proche d'un produit professionnel."
    ], table=(["Phase", "Objectif", "Livrables"], [
        ("Phase 1", "Securiser l'acces", "JWT, RBAC, audit logs"),
        ("Phase 2", "Temps reel", "WebSockets, live metrics"),
        ("Phase 3", "Monitoring avance", "Prometheus, Grafana"),
        ("Phase 4", "Cloud multi-ressources", "VM multiples, storage, network security groups"),
        ("Phase 5", "CI/CD", "GitHub Actions, tests, scans securite"),
    ], [1.0, 2.3, 3.2]))

    section_page(doc, "45. Conclusion", [
        "Cloud Admin Platform est devenu une plateforme DevOps complete permettant de superviser et administrer une infrastructure Cloud reelle.",
        "Le projet s'appuie sur React, FastAPI, Docker, Azure VM, Nginx, DuckDNS, Let's Encrypt et Vercel. Il montre une capacite a concevoir, deployer, diagnostiquer et ameliorer une application Full Stack moderne.",
        "Au-dela de son etat actuel, le projet est extensible vers une solution professionnelle avec authentification, roles, monitoring avance, temps reel et pipeline CI/CD."
    ])

    h1(doc, "Annexe A - Extrait des endpoints")
    add_table(doc, ["Methode", "Endpoint", "Fonction", "Source"], ENDPOINTS, [0.75, 1.8, 2.35, 1.6])
    doc.add_page_break()

    h1(doc, "Annexe B - Exemples de blocs de code")
    h2(doc, "Configuration Axios")
    add_code(doc, """
const API_URL = process.env.REACT_APP_API_URL || 'http://localhost:8000';
const api = axios.create({
  baseURL: API_URL,
  timeout: 10000,
  headers: { 'Content-Type': 'application/json' }
});
""")
    h2(doc, "Docker Compose")
    add_code(doc, """
services:
  postgres:
    image: postgres:16
  backend:
    build: ./backend
    ports:
      - "8000:8000"
    volumes:
      - /var/run/docker.sock:/var/run/docker.sock
      - ./:/workspace
""")
    h2(doc, "Theme automatique")
    add_code(doc, """
const getAutoTheme = () => {
  const hour = new Date().getHours();
  const isNight = hour >= 19 || hour < 7;
  const prefersDark = window.matchMedia('(prefers-color-scheme: dark)').matches;
  return isNight || prefersDark ? 'dark' : 'light';
};
""")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
