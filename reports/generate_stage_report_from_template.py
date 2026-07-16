from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_cloud_admin_report import ASSET_DIR, ENDPOINTS, SCREENSHOTS, TECH, make_diagrams


ROOT = Path(r"D:\Projet Stage Deweb\New projet\Projects\cloud-admin-platform")
TEMPLATE = Path(r"C:\Users\yassi\Downloads\Modèle de rapport de  stage INGENIEUR D ETUDES.docx")
OUT = ROOT / "reports" / "Cloud_Admin_Platform_Rapport_Stage_Modele_Final.docx"


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_after_cover(doc):
    body = doc._body._element
    children = list(body)
    keep = []
    for child in children:
        if child.tag.endswith("tbl"):
            keep.append(child)
            break
    sect = children[-1] if children and children[-1].tag.endswith("sectPr") else None
    for child in children:
        if child in keep or child is sect:
            continue
        remove_element(child)


def set_run(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, text="", style="Body Text", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_title(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper() if level == 1 else text)
    set_run(r, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run("· " + item)
        set_run(r)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run(r, size=size, bold=bold)


def fill_cover(doc):
    table = doc.tables[0]
    set_cell_text(table.cell(0, 0), "\n\nCloudAdmin\nCloud Admin Platform", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        table.cell(1, 0),
        "\n\nCYCLE INGENIEUR DE L'ECOLE POLYTECHNIQUE D'AGADIR\nFILIERE INGENIERIE INFORMATIQUE / CLOUD & DEVOPS",
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(table.cell(2, 0), "STAGE INGENIEUR D'ETUDES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        table.cell(3, 0),
        "Cloud Admin Platform\nPlateforme Full Stack de supervision et d'administration Cloud\n\n\n\n\n\n\nRéalisé par :\tSuperviseur :\nYassine Rachid\tM/Mme Nom Prénom\n\tTuteur :\nDr. El Mehdi KIBBOU\n\n\n\n\nAnnée universitaire 2025/2026",
        size=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def add_toc_placeholder(doc):
    add_title(doc, "TABLE DES MATIERES", 1)
    add_p(doc, "La table des matières automatique est configurée pour être mise à jour à l'ouverture du document dans Microsoft Word.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Cliquez avec le bouton droit puis choisissez Mettre à jour le champ."
    run.append(text)
    fld.append(run)
    p._p.append(fld)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Tableau Grille"
    except Exception:
        pass
    table.autofit = True
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=9)
    doc.add_paragraph()
    return table


def add_picture_if_exists(doc, path, caption):
    p = Path(path)
    if not p.exists():
        return
    try:
        doc.add_picture(str(p), width=Inches(6.3))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    except Exception:
        pass


def add_extended_study(doc):
    add_title(doc, "Analyse détaillée du besoin", 2)
    needs = [
        ("Centralisation", "Avant la réalisation de CloudAdmin, les informations d'administration étaient dispersées entre plusieurs outils : terminal Linux, Docker CLI, Azure CLI, fichiers de logs, règles UFW et état du dépôt Git. Cette dispersion rend le diagnostic plus lent, surtout lorsqu'un administrateur doit vérifier rapidement l'état global d'une plateforme. La centralisation dans une interface Web réduit le temps de consultation et facilite la lecture des informations essentielles."),
        ("Données réelles", "Un point important du projet est l'exigence d'afficher des données réelles. Une plateforme d'administration perd de sa valeur si les cartes, tableaux ou graphiques reposent sur des données statiques. Le backend a donc été conçu pour interroger directement le système, les conteneurs Docker, Git, Terraform, le réseau et le stockage. Cette approche permet à l'interface de représenter l'état réel de l'environnement."),
        ("Actions", "La supervision seule n'est pas suffisante pour un produit d'administration. L'utilisateur doit pouvoir agir sur l'infrastructure, dans un cadre contrôlé. Les boutons Start, Stop, Restart, Reload, Fetch, Pull ou Push représentent des points d'action concrets. Leur connexion au backend transforme l'application en outil opérationnel et non en simple tableau de bord visuel."),
        ("Expérience", "L'expérience utilisateur constitue un objectif important du projet. L'interface doit inspirer confiance, rester lisible en dark mode comme en light mode et donner une impression de produit SaaS moderne. Les animations, les transitions et les micro-interactions ne sont pas uniquement décoratives : elles confirment les actions, guident le regard et rendent l'application plus agréable à utiliser."),
    ]
    for title, text in needs:
        add_title(doc, title, 3)
        add_p(doc, text)


def add_design_details(doc):
    add_title(doc, "Conception fonctionnelle", 2)
    modules = [
        ("Dashboard", "Le dashboard fournit une synthèse immédiate : charge CPU, mémoire, disque, état Docker, API, base de données, versions des outils et santé générale. Cette page sert de point d'entrée pour identifier rapidement les anomalies."),
        ("Virtual Machines", "La page Virtual Machines expose l'état de la machine Cloud, sa localisation, ses adresses IP et les actions de cycle de vie. Elle doit rester synchronisée avec les données Azure ou avec le fallback local selon l'environnement."),
        ("Docker", "La page Docker affiche les conteneurs actifs, leurs images, ports, états, réseaux et actions. Les boutons de démarrage, arrêt et redémarrage sont essentiels pour démontrer une administration concrète."),
        ("Networks", "La page Networks lit les interfaces réseau, les adresses IP, le MTU, le statut de lien, la passerelle, les DNS et les compteurs RX/TX. Elle aide au diagnostic de connectivité."),
        ("Storage", "La page Storage analyse l'espace disque, les volumes Docker, le montage racine et l'utilisation globale. Elle permet de détecter les risques liés à la saturation du stockage."),
        ("Metrics", "La page Metrics regroupe les historiques CPU, RAM, swap, disque, réseau et load average. Les graphiques progressifs améliorent la compréhension visuelle de l'évolution des ressources."),
        ("Logs", "La page Logs centralise les traces utiles. Elle doit gérer le filtrage, la recherche, le responsive et l'affichage de grands volumes sans casser le conteneur."),
        ("Security", "Les pages Firewall, SSH Keys et IAM exposent les informations critiques de sécurité : règles UFW, clés autorisées, comptes utilisateurs et shells associés."),
        ("DevOps", "Les pages GitHub, Terraform et Docker Compose connectent le produit aux pratiques DevOps : état Git, commits, remote, workspace Terraform, ressources et stack Compose."),
    ]
    for name, desc in modules:
        add_title(doc, name, 3)
        add_p(doc, desc)
        add_p(doc, "Pour ce module, la conception respecte trois principes : une API dédiée, un composant React isolé et une interface capable d'afficher les états de chargement, d'erreur, de vide et de succès. Cette discipline rend l'application plus maintenable et évite les comportements incohérents entre les pages.")


def add_backend_details(doc):
    add_title(doc, "Conception backend et intégration système", 2)
    topics = [
        ("FastAPI", "FastAPI a été retenu pour sa rapidité, sa clarté et sa compatibilité avec Pydantic. Les endpoints REST sont simples à tester et peuvent être documentés automatiquement. Cette technologie convient bien à une application qui doit exposer des informations système sous forme JSON."),
        ("psutil", "La bibliothèque psutil fournit un accès fiable aux métriques CPU, mémoire, disque, swap, réseau et processus. Elle permet d'éviter des parsings fragiles de commandes shell lorsque les données sont disponibles par API Python."),
        ("Docker SDK", "Le Docker SDK permet de lister les conteneurs, lire leur état, récupérer les logs et déclencher des actions. Il est utilisé en complément de la CLI pour rendre les opérations plus structurées."),
        ("subprocess", "Certaines informations nécessitent encore l'appel de commandes externes, par exemple Git, Terraform, UFW ou Azure CLI. Ces appels sont encapsulés côté backend pour éviter d'exposer directement le terminal au frontend."),
        ("Gestion des erreurs", "Les endpoints doivent distinguer une absence réelle de données, une dépendance non installée, une permission insuffisante et une erreur d'exécution. Cette distinction améliore la qualité des messages affichés dans l'interface."),
    ]
    for title, text in topics:
        add_title(doc, title, 3)
        add_p(doc, text)
        add_p(doc, "Dans une évolution professionnelle, chaque endpoint pourrait ajouter un code d'erreur fonctionnel, un identifiant de corrélation et une trace d'audit afin de faciliter le support et la sécurité.")


def add_frontend_details(doc):
    add_title(doc, "Conception frontend et design system", 2)
    add_p(doc, "Le frontend est construit comme une application React à pages spécialisées. Chaque page consomme un service API, transforme les réponses en données d'affichage et propose des composants réutilisables comme les cartes statistiques, les tableaux, les états vides, les boutons d'action et les notifications.")
    add_p(doc, "Le design system repose sur des variables CSS afin de garantir la cohérence entre dark mode et light mode. Les couleurs codées en dur ont été évitées pour réduire les problèmes de contraste et faciliter l'évolution de l'identité visuelle.")
    add_p(doc, "Les animations ont été pensées comme des micro-feedbacks : apparition progressive des pages, léger flottement des cartes, survol plus expressif, boutons plus réactifs et arrière-plan vivant. Cette stratégie permet de dynamiser l'application sans nuire à la lisibilité.")
    add_p(doc, "Le mode automatique constitue un élément différenciant : au premier accès, l'utilisateur choisit dark, light ou automatique. Ensuite, sa préférence reste enregistrée. En mode automatique, l'application tient compte de l'appareil et du moment de la journée.")
    add_table(doc, ["Composant UX", "Rôle", "Impact"], [
        ("Toast", "Informer après une action", "Réduit l'incertitude utilisateur"),
        ("Skeleton", "Afficher un chargement propre", "Evite les écrans vides"),
        ("Confirm dialog", "Protéger les actions sensibles", "Diminue les erreurs"),
        ("Hover", "Signaler l'interactivité", "Rend l'interface plus vivante"),
        ("Transition", "Accompagner le changement de page", "Améliore la fluidité"),
        ("Theme prompt", "Personnaliser dès le premier accès", "Renforce le confort"),
    ])


def add_tests_and_validation(doc):
    add_title(doc, "Tests, validation et audit", 2)
    checks = [
        ("Validation API", "Chaque page importante a été vérifiée afin de confirmer qu'elle interroge bien son endpoint. Les réponses JSON ont été comparées avec l'affichage afin d'éviter les erreurs de mapping."),
        ("Validation UI", "Les pages ont été vérifiées en dark mode et en light mode. Les problèmes de texte noir sur fond sombre et de fond noir en light mode ont été corrigés avec des design tokens."),
        ("Validation responsive", "Les tableaux et conteneurs ont été analysés pour éviter les dépassements. La page Logs a été identifiée comme critique car elle affiche des lignes longues et doit rester lisible."),
        ("Validation actions", "Les boutons Start, Stop, Restart et Reload ont été reliés aux appels backend. Après une action, l'interface recharge les données pour afficher l'état actualisé."),
        ("Validation GitHub", "La page GitHub a été simplifiée en supprimant le bloc Commit changes et en ajoutant des informations plus utiles comme l'URL du repository, le nombre de commits et la taille du dépôt."),
    ]
    for title, text in checks:
        add_title(doc, title, 3)
        add_p(doc, text)
    add_table(doc, ["Test", "Méthode", "Résultat attendu"], [
        ("Chargement Dashboard", "GET /api/dashboard", "Cartes remplies avec données système"),
        ("Docker actions", "POST start/stop/restart", "Etat conteneur mis à jour"),
        ("Light mode", "Basculer thème", "Fond clair sans noir dominant"),
        ("Logs", "Filtrer et rechercher", "Tableau lisible sans débordement"),
        ("GitHub", "GET /api/github", "Repository details complets"),
        ("Settings", "Afficher API config", "Bloc droit non compressé"),
    ])


def add_risks_and_improvements(doc):
    add_title(doc, "Risques, limites et pistes d'amélioration", 2)
    add_p(doc, "Même si la plateforme atteint ses objectifs principaux, certaines limites doivent être prises en compte avant une mise en production complète. La première concerne la sécurité des actions : démarrer, arrêter ou modifier une ressource doit être réservé à des utilisateurs authentifiés et autorisés.")
    add_p(doc, "La deuxième limite concerne l'exposition du socket Docker. Le montage de /var/run/docker.sock donne un pouvoir important au backend. En production, il faut restreindre ce privilège, isoler les permissions et ajouter un audit précis des actions.")
    add_p(doc, "La troisième limite concerne le temps réel. Les métriques sont actuellement récupérées par appels API et rafraîchissements. Une architecture WebSocket améliorerait le suivi continu et réduirait la latence perçue.")
    add_table(doc, ["Risque", "Impact", "Mesure proposée"], [
        ("Absence d'authentification", "Accès non autorisé", "JWT et RBAC"),
        ("Docker socket exposé", "Contrôle élevé de l'hôte", "Proxy Docker sécurisé"),
        ("Commandes système", "Erreurs ou permissions", "Validation et sandboxing"),
        ("Données volumineuses", "Tables lentes", "Pagination et virtualisation"),
        ("Secrets", "Fuite de configuration", "Variables d'environnement et coffre secret"),
        ("Monitoring limité", "Diagnostic incomplet", "Prometheus, Grafana, alerting"),
    ])


def add_technical_dossier(doc):
    add_title(doc, "Dossier technique détaillé", 2)
    domains = [
        ("Architecture globale", "L'architecture globale de CloudAdmin suit une logique en couches. Le navigateur exécute l'application React, le frontend consomme les endpoints REST, le reverse proxy Nginx expose l'API en HTTPS et le backend FastAPI interroge les composants système. Cette organisation facilite le diagnostic car chaque responsabilité est localisée dans une couche précise."),
        ("Modélisation des données", "Les données manipulées par l'application ne proviennent pas d'un modèle métier classique mais de ressources techniques : conteneurs, interfaces réseau, volumes, métriques, utilisateurs, règles firewall et informations Git. Cette particularité impose de normaliser les réponses afin que l'interface puisse afficher des états homogènes malgré la diversité des sources."),
        ("Backend FastAPI", "Le backend agit comme une passerelle contrôlée entre le frontend et l'environnement système. Il protège l'utilisateur de la complexité des commandes bas niveau et transforme les résultats en JSON. Cette transformation améliore la lisibilité, réduit la duplication côté frontend et rend les pages plus simples à maintenir."),
        ("Frontend React", "Le frontend a été conçu pour donner une expérience de console SaaS. Les composants réutilisables permettent de conserver la même logique d'affichage sur les pages : cartes statistiques, tableaux, badges, boutons d'action, états vides et messages d'erreur. Cette cohérence réduit la charge cognitive de l'utilisateur."),
        ("Gestion du thème", "Le thème constitue un élément important de confort. Le dark mode est adapté aux usages nocturnes et aux environnements techniques, tandis que le light mode améliore la lisibilité en journée. Le mode automatique permet de respecter la préférence système et le moment de la journée sans demander une action répétée à l'utilisateur."),
        ("Docker et conteneurisation", "Docker simplifie l'exécution du backend et de la base PostgreSQL. Le montage du socket Docker donne au backend la capacité de superviser les conteneurs de l'hôte. Cette décision est puissante mais doit être accompagnée de contrôles stricts si l'application est exposée publiquement."),
        ("Azure et hébergement", "Azure VM sert d'environnement cible pour le backend et les outils système. Cette approche permet de travailler sur une infrastructure réelle, avec un système Linux, Docker, Git, Terraform et des configurations réseau concrètes. Elle rend le projet plus crédible qu'une simple maquette locale."),
        ("Nginx et HTTPS", "Nginx joue le rôle de reverse proxy entre Internet et l'API FastAPI. Il permet de centraliser l'exposition, de forcer HTTPS et de préparer des règles de sécurité supplémentaires. Let's Encrypt et DuckDNS assurent un accès public sécurisé avec un domaine identifiable."),
        ("GitHub et traçabilité", "GitHub permet de suivre l'évolution du projet, de documenter les commits et de préparer un workflow de déploiement. La page GitHub de CloudAdmin expose des informations utiles comme la branche courante, l'état de modification, le remote, le nombre de commits et l'historique récent."),
        ("Terraform et infrastructure as code", "Terraform représente la dimension Infrastructure as Code du projet. Même si certaines opérations dépendent de la disponibilité locale de Terraform, l'intégration prépare une évolution vers un pilotage plus complet des ressources Cloud depuis l'interface."),
        ("Logs et observabilité", "Les logs sont essentiels pour comprendre le comportement d'une application distribuée. La plateforme affiche les traces backend, Docker et système afin de réduire le temps de diagnostic. Le filtrage par niveau et la recherche textuelle améliorent l'exploitation quotidienne."),
        ("Sécurité opérationnelle", "La sécurité concerne les accès, les actions, les secrets, les permissions et l'audit. Le projet pose les bases de cette sécurité avec les pages Firewall, SSH Keys et IAM, mais il doit évoluer vers une authentification forte, une gestion des rôles et des journaux d'audit complets."),
    ]
    focus = [
        "Objectif fonctionnel",
        "Choix technique",
        "Implémentation",
        "Difficultés rencontrées",
        "Résultat obtenu",
        "Amélioration possible",
    ]
    for domain, base in domains:
        add_title(doc, domain, 3)
        add_p(doc, base)
        for item in focus:
            add_p(
                doc,
                f"{item} : dans le contexte de {domain.lower()}, le travail a consisté à relier une exigence utilisateur à une solution technique vérifiable. La priorité a été de conserver une architecture compréhensible, de limiter les données simulées et de rendre l'information exploitable depuis l'interface. Cette démarche a permis de transformer le projet en plateforme démontrable, utile pour un PFE, un portfolio GitHub et un entretien technique.",
            )


def add_operational_manual(doc):
    add_title(doc, "Manuel d'exploitation", 2)
    scenarios = [
        ("Démarrage local", "Le développeur lance le backend FastAPI en mode reload puis démarre le frontend React. Cette configuration permet de tester rapidement les modifications d'interface et de vérifier les réponses API sur localhost."),
        ("Démarrage Docker", "L'administrateur utilise docker compose up -d --build pour reconstruire et relancer le backend et PostgreSQL. Cette méthode correspond davantage à un environnement de production ou de préproduction."),
        ("Diagnostic Docker", "En cas de conteneur arrêté, l'utilisateur consulte la page Docker, lit l'état, vérifie les logs et déclenche si nécessaire une action restart. Le retour visuel confirme ensuite l'état actualisé."),
        ("Diagnostic réseau", "La page Networks permet de vérifier les interfaces, les IP, le MTU, la passerelle et les DNS. Ces informations sont utiles lorsqu'une API est inaccessible ou lorsqu'un service ne communique pas correctement."),
        ("Diagnostic stockage", "La page Storage signale les volumes, le disque racine et l'espace disponible. Un manque de stockage peut provoquer des erreurs Docker, PostgreSQL ou système."),
        ("Diagnostic logs", "La page Logs aide à rechercher les erreurs, avertissements et traces d'accès. Le filtre par niveau simplifie la lecture lorsque le volume de logs augmente."),
        ("Diagnostic Git", "La page GitHub permet de savoir si le dépôt est modifié, si la branche est en retard ou si un push est nécessaire. Elle complète le workflow de développement et de déploiement."),
        ("Diagnostic Terraform", "La page Terraform vérifie la présence de Terraform, les fichiers disponibles, le workspace et les ressources connues. Elle prépare l'intégration d'opérations IaC plus avancées."),
    ]
    for title, text in scenarios:
        add_title(doc, title, 3)
        add_p(doc, text)
        add_p(doc, "La procédure recommandée consiste à observer l'état, identifier la source de l'anomalie, déclencher une action limitée et vérifier le résultat après rafraîchissement. Cette démarche évite les décisions hâtives et favorise une exploitation maîtrisée.")
        add_p(doc, "Dans une version future, chaque scénario pourrait être accompagné d'un journal d'audit, d'une demande de confirmation renforcée et d'une notification temps réel envoyée aux administrateurs concernés.")


def add_project_management(doc):
    add_title(doc, "Gestion du projet technique", 2)
    add_p(doc, "La gestion du projet a suivi une approche itérative. Les premières étapes ont porté sur l'identification des pages incomplètes, puis sur la correction des appels API, l'amélioration de l'interface et la stabilisation du déploiement. Chaque itération a résolu un problème concret observé dans l'application.")
    add_table(doc, ["Phase", "Objectif", "Livrable"], [
        ("Analyse", "Identifier les pages statiques et les N/A", "Liste des corrections"),
        ("Backend", "Compléter les endpoints", "API connectée"),
        ("Frontend", "Mapper les données réelles", "Pages dynamiques"),
        ("UX", "Ajouter thème et animations", "Interface plus professionnelle"),
        ("Documentation", "Créer README et rapport", "Livrables GitHub/PFE"),
        ("Déploiement", "Pousser sur GitHub", "Version accessible"),
    ])
    for title in ["Suivi des besoins", "Priorisation", "Validation", "Documentation", "Capitalisation"]:
        add_title(doc, title, 3)
        add_p(doc, f"{title} : cette étape a permis de transformer les remarques fonctionnelles en tâches techniques concrètes. Les corrections ont été priorisées selon leur impact sur la démonstration du projet : données réelles, actions fonctionnelles, lisibilité, cohérence graphique et crédibilité du dépôt GitHub.")
        add_p(doc, "Cette méthode a aussi permis de conserver un lien direct entre les captures d'écran, les problèmes observés et les corrections réalisées dans le code.")


def build_report():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_diagrams()
    arch = ASSET_DIR / "architecture.png"
    flow = ASSET_DIR / "dataflow.png"
    dep = ASSET_DIR / "deployment.png"

    doc = Document(str(TEMPLATE))
    fill_cover(doc)
    clear_after_cover(doc)

    page_break(doc)
    add_title(doc, "REMERCIEMENTS", 1)
    add_p(doc, "Je tiens à exprimer ma sincère reconnaissance à l'ensemble des personnes qui ont contribué à la réalisation de ce projet de fin d'études, intitulé Cloud Admin Platform.")
    add_p(doc, "Je remercie particulièrement mon encadrant pédagogique pour son accompagnement, ses orientations méthodologiques et ses remarques constructives. Je remercie également les personnes qui m'ont aidé à clarifier les besoins techniques liés à l'administration Cloud, à Docker, à Azure, au monitoring et aux pratiques DevOps.")
    add_p(doc, "Ce projet a représenté une occasion importante de consolider mes compétences en développement full stack, en administration Linux, en intégration API, en déploiement Docker et en mise en production d'une architecture sécurisée.")

    page_break(doc)
    add_title(doc, "RESUME", 1)
    add_p(doc, "Cloud Admin Platform est une plateforme Web Full Stack de supervision et d'administration d'une infrastructure Cloud réelle. Elle permet de centraliser, dans une interface moderne, les informations liées aux machines virtuelles, aux conteneurs Docker, aux réseaux, au stockage, aux métriques système, aux logs, aux règles firewall, aux utilisateurs Linux, aux clés SSH, à GitHub, à Terraform et à Docker Compose.")
    add_p(doc, "Le frontend est développé avec React et communique avec un backend FastAPI par API REST. Le backend collecte les données directement depuis le système Linux, Docker, Git, Terraform, UFW, Azure CLI et psutil. L'objectif principal est d'éviter les données statiques et de garantir que les informations affichées reflètent l'état réel de l'environnement.")
    add_p(doc, "Mots clés : Cloud, DevOps, React, FastAPI, Docker, Azure, Monitoring, Terraform, Nginx, API REST.")

    add_title(doc, "ABSTRACT", 1)
    add_p(doc, "Cloud Admin Platform is a full-stack Web platform designed to monitor and administrate a real Cloud infrastructure. It centralizes virtual machines, Docker containers, networks, storage, system metrics, logs, firewall rules, Linux users, SSH keys, GitHub, Terraform and Docker Compose in a modern web interface.")
    add_p(doc, "The frontend is built with React and consumes REST APIs exposed by a FastAPI backend. The backend retrieves live data from Linux, Docker, Git, Terraform, UFW, Azure CLI and psutil. The main objective is to avoid static business data and provide a reliable operational dashboard based on real infrastructure sources.")
    add_p(doc, "Key words: Cloud, DevOps, React, FastAPI, Docker, Azure, Monitoring, Terraform, Nginx, REST API.")

    page_break(doc)
    add_toc_placeholder(doc)

    page_break(doc)
    add_title(doc, "TABLE DES ILLUSTRATIONS", 1)
    illustrations = [
        ("Figure 1", "Architecture générale de CloudAdmin"),
        ("Figure 2", "Flux de données Frontend - Backend - Système"),
        ("Figure 3", "Pipeline de déploiement Azure / Vercel"),
        ("Figure 4", "Interface Networks en light mode"),
        ("Figure 5", "Inventaire Docker avec actions"),
    ]
    add_table(doc, ["Référence", "Légende"], illustrations)
    add_title(doc, "LISTE DES TABLEAUX", 1)
    add_table(doc, ["Référence", "Titre"], [("Tableau 1", "Technologies utilisées"), ("Tableau 2", "Endpoints REST"), ("Tableau 3", "Risques et solutions")])

    page_break(doc)
    add_title(doc, "INTRODUCTION GENERALE", 1)
    intro_paras = [
        "La transformation numérique des organisations pousse les équipes techniques à rechercher des solutions capables de superviser, diagnostiquer et administrer rapidement leurs environnements Cloud. Dans ce contexte, les plateformes DevOps jouent un rôle essentiel, car elles rapprochent les informations système, les outils de déploiement, les métriques applicatives et les actions d'exploitation.",
        "Le projet Cloud Admin Platform s'inscrit dans cette logique. Il vise à construire une plateforme Web moderne qui donne à un administrateur une vision centralisée de son infrastructure. Le projet couvre le frontend, le backend, l'intégration système, Docker, Git, Terraform, les logs, la sécurité de base et le déploiement Cloud.",
        "Le stage a donc porté sur la conception, le développement et l'amélioration d'une solution Full Stack connectée à une infrastructure réelle. Le travail a inclus la correction des pages sans données, la connexion complète aux API, l'amélioration de l'expérience utilisateur, l'ajout du mode sombre et clair automatique, ainsi que la création d'un design plus professionnel.",
        "Le présent rapport est organisé en trois chapitres. Le premier présente le contexte, l'organisme d'accueil et le besoin fonctionnel. Le deuxième expose la recherche documentaire et les technologies utilisées. Le troisième détaille les missions réalisées, l'architecture, les développements, les corrections, les résultats obtenus et les perspectives d'amélioration.",
    ]
    for para in intro_paras:
        add_p(doc, para)

    page_break(doc)
    add_p(doc, "CHAPITRE 1 :", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    add_p(doc, "Présentation de l'entreprise et du contexte du projet", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    page_break(doc)
    add_p(doc, "Chapitre 1 : Présentation de l'entreprise et du contexte du projet", style="Body Text", bold=True)
    add_title(doc, "INTRODUCTION", 2)
    add_p(doc, "Ce chapitre présente le cadre général du projet Cloud Admin Platform, son contexte technique et les besoins auxquels il répond. L'objectif est de situer le projet dans un environnement DevOps et Cloud moderne, où la supervision, l'automatisation et la fiabilité des données sont devenues indispensables.")
    add_title(doc, "RAISON SOCIALE ET STATUT JURIDIQUE", 2)
    add_p(doc, "Le projet est réalisé dans un cadre académique et professionnel de type Projet de Fin d'Etudes. Il peut être présenté comme un portfolio DevOps et Cloud Administration démontrant des compétences concrètes en développement Web, en administration système, en conteneurisation et en déploiement Cloud.")
    add_title(doc, "ACTIVITES", 3)
    add_p(doc, "L'activité principale du projet consiste à fournir une interface d'administration Cloud capable de récupérer les informations depuis des sources réelles. La plateforme regroupe les données liées aux machines virtuelles, aux conteneurs, au réseau, au stockage, aux métriques, aux logs et aux outils DevOps.")
    add_title(doc, "ORGANISATION DU PROJET", 3)
    add_p(doc, "L'organisation technique repose sur trois parties principales : une interface React, une API FastAPI et une infrastructure d'exécution basée sur Azure, Docker Compose, Nginx, DuckDNS et Let's Encrypt. Cette séparation facilite la maintenance, le déploiement et l'évolution du produit.")
    add_bullets(doc, ["Frontend React pour l'expérience utilisateur.", "Backend FastAPI pour l'accès aux données et actions système.", "Infrastructure Azure VM pour l'hébergement de l'API et des outils DevOps.", "Déploiement frontend sur Vercel avec variables d'environnement.", "Reverse proxy Nginx et HTTPS pour sécuriser les échanges."])
    add_title(doc, "SYSTEME DE SERVICE", 3)
    add_p(doc, "Le service rendu par CloudAdmin est une console Web de supervision et d'administration. L'utilisateur consulte les pages métiers, déclenche certaines actions comme démarrer ou arrêter un conteneur, recharge les données et vérifie les états système en temps réel ou quasi réel.")
    add_extended_study(doc)
    add_title(doc, "LES METIERS DU TECHNICIEN AU SEIN DU PROJET", 2)
    add_p(doc, "Le projet mobilise des compétences proches de celles d'un technicien ou ingénieur Cloud/DevOps : analyse d'incident, surveillance des ressources, gestion Docker, vérification réseau, gestion Git, lecture des logs et sécurisation progressive d'une infrastructure.")
    add_title(doc, "CONCLUSION", 2)
    add_p(doc, "Ce premier chapitre montre que CloudAdmin répond à un besoin opérationnel clair : rassembler les informations d'administration dans une interface unique, lisible et connectée à l'état réel de l'infrastructure.")

    page_break(doc)
    add_p(doc, "CHAPITRE 2 :", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    add_p(doc, "Recherche documentaire", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    page_break(doc)
    add_p(doc, "Chapitre 2 : Recherche documentaire", style="Body Text", bold=True)
    add_title(doc, "INTRODUCTION", 2)
    add_p(doc, "Ce chapitre présente les outils, méthodes et technologies utilisés pour concevoir CloudAdmin. La recherche documentaire a permis d'identifier les bonnes pratiques liées aux architectures Full Stack, à l'administration Linux, à Docker, au monitoring et au déploiement sécurisé.")
    add_title(doc, "OUTILS ET METHODES UTILISEES", 2)
    add_table(doc, ["Domaine", "Technologie", "Rôle"], TECH[:18])
    add_design_details(doc)
    add_title(doc, "ARCHITECTURE TECHNIQUE", 2)
    add_p(doc, "L'architecture retenue sépare clairement la couche présentation, la couche API et la couche système. Cette séparation permet de maintenir une interface évolutive tout en isolant les opérations sensibles dans le backend.")
    add_picture_if_exists(doc, arch, "Figure 1 : Architecture générale de CloudAdmin.")
    add_picture_if_exists(doc, flow, "Figure 2 : Flux de données entre le frontend, le backend et les sources système.")
    add_title(doc, "API REST ET SOURCES DE DONNEES", 2)
    add_p(doc, "Chaque page de la plateforme consomme un endpoint spécifique. Les informations proviennent de commandes système, de bibliothèques Python ou d'outils CLI. L'objectif est d'obtenir une donnée réelle et vérifiable, et non une donnée simulée.")
    add_table(doc, ["Méthode", "Endpoint", "Description", "Source"], ENDPOINTS[:18])
    add_table(doc, ["Méthode", "Endpoint", "Description", "Source"], ENDPOINTS[18:])
    add_backend_details(doc)
    add_title(doc, "DEPLOIEMENT ET SECURISATION", 2)
    add_p(doc, "Le backend est prévu pour fonctionner sur une machine virtuelle Azure via Docker Compose. Le frontend peut être déployé sur Vercel et consommer l'API HTTPS exposée par Nginx. DuckDNS fournit un nom de domaine et Let's Encrypt permet d'activer le chiffrement TLS.")
    add_picture_if_exists(doc, dep, "Figure 3 : Pipeline de déploiement et d'hébergement.")
    add_title(doc, "CONCLUSION", 2)
    add_p(doc, "La recherche documentaire a permis de retenir une pile technologique cohérente avec les objectifs du projet : React pour l'interface, FastAPI pour l'API, Docker pour l'exécution, Azure pour l'hébergement et Nginx/HTTPS pour l'exposition sécurisée.")

    page_break(doc)
    add_p(doc, "CHAPITRE 3 :", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    add_p(doc, "Mission et travaux réalisés", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    page_break(doc)
    add_p(doc, "Chapitre 3 : Missions et travaux réalisés", style="Body Text", bold=True)
    add_title(doc, "INTRODUCTION", 2)
    add_p(doc, "Ce chapitre détaille les travaux réalisés pendant le projet : connexion aux API, correction des pages vides, amélioration des actions d'administration, optimisation du light mode, enrichissement du dark mode, animations, génération du rapport et préparation du dépôt GitHub.")
    add_title(doc, "LES MISSIONS EFFECTUEES", 2)
    add_title(doc, "Mission principale : réalisation de la plateforme CloudAdmin", 3)
    add_p(doc, "La mission principale consistait à transformer une interface de supervision statique ou partiellement connectée en une plateforme cohérente, dynamique et branchée sur des API réelles. Plusieurs pages présentaient des valeurs N/A ou des états vides alors que les données existaient côté backend.")
    add_bullets(doc, ["Audit des pages sans données.", "Correction de la récupération API via les services frontend.", "Connexion des boutons Start, Stop, Restart et Reload.", "Suppression des données métiers statiques.", "Correction du tableau des logs et des débordements responsive.", "Amélioration du light mode et suppression des fonds noirs inadaptés.", "Ajout de micro-interactions et d'animations en dark et light mode.", "Mise en place d'une préférence de thème persistante par utilisateur."])
    add_title(doc, "Correction des données N/A", 3)
    add_p(doc, "Les valeurs N/A apparaissaient lorsque les composants frontend ne consommaient pas correctement les réponses API ou lorsque le mapping des champs backend ne couvrait pas toutes les variantes de données. La correction a consisté à harmoniser les structures de réponses et à utiliser des valeurs de fallback uniquement lorsque la donnée était réellement absente.")
    add_title(doc, "Actions d'administration", 3)
    add_p(doc, "Les boutons d'action des pages Virtual Machines, Docker, Terraform et GitHub ont été reliés à des endpoints backend dédiés. L'utilisateur peut ainsi déclencher des opérations concrètes comme le démarrage, l'arrêt ou le redémarrage d'un service, puis voir l'interface se mettre à jour.")
    add_title(doc, "Expérience utilisateur et thème", 3)
    add_p(doc, "Le système de thème a été enrichi avec trois modes : light, dark et automatique. Au premier accès, l'utilisateur peut choisir sa préférence. Le choix est enregistré localement afin de rester actif lors des visites suivantes. Le mode automatique tient compte de la préférence système et du cycle jour/nuit.")
    add_title(doc, "Animations et micro-interactions", 3)
    add_p(doc, "Les animations ont été ajoutées pour donner une impression plus professionnelle : transitions de pages, effets de survol sur les cartes, animation du fond, apparition progressive des graphiques, feedback visuel sur les boutons et notifications après les actions.")
    add_frontend_details(doc)
    add_tests_and_validation(doc)
    add_technical_dossier(doc)
    add_operational_manual(doc)
    add_project_management(doc)
    add_title(doc, "Captures d'écran de l'application", 2)
    for idx, (path, caption) in enumerate(SCREENSHOTS, 4):
        add_picture_if_exists(doc, path, f"Figure {idx} : {caption}")
    add_title(doc, "Résultats obtenus", 2)
    add_p(doc, "Le résultat est une plateforme plus proche d'un produit SaaS réel : les pages sont connectées, les actions principales fonctionnent, le thème est plus cohérent, les animations améliorent la perception de qualité et le dépôt GitHub dispose d'une documentation professionnelle.")
    add_table(doc, ["Problème initial", "Correction réalisée", "Résultat"], [
        ("Pages sans données", "Connexion API et mapping dynamique", "Données réelles affichées"),
        ("N/A sur plusieurs pages", "Fallback contrôlé et endpoints enrichis", "Valeurs cohérentes"),
        ("Boutons non fonctionnels", "Actions POST backend", "Start/Stop/Restart opérationnels"),
        ("Light mode trop sombre", "Palette claire et fond adapté", "Interface plus lisible"),
        ("Animations faibles", "Framer Motion et CSS transitions", "Expérience plus dynamique"),
        ("Branding doublonné", "Suppression du doublon navbar", "Logo CloudAdmin clair"),
    ])
    add_risks_and_improvements(doc)
    add_title(doc, "CONCLUSION", 2)
    add_p(doc, "Les missions réalisées ont permis de finaliser un produit présentable, robuste et connecté à une infrastructure réelle. Les améliorations apportées renforcent à la fois la valeur technique et la valeur visuelle du projet.")

    page_break(doc)
    add_title(doc, "CONCLUSION GENERALE", 1)
    add_p(doc, "Cloud Admin Platform est devenu une plateforme DevOps complète permettant de superviser et d'administrer une infrastructure Cloud réelle depuis une interface Web moderne. Le projet combine React, FastAPI, Docker, Azure, Nginx, DuckDNS, GitHub et Terraform dans une architecture cohérente.")
    add_p(doc, "Sur le plan technique, le projet démontre la capacité à concevoir des API REST, à intégrer des sources système, à automatiser des actions d'administration et à gérer un déploiement Cloud sécurisé. Sur le plan frontend, il montre une attention particulière au design, au responsive, au dark/light mode et aux micro-interactions.")
    add_p(doc, "Les pistes d'amélioration concernent l'authentification JWT, la gestion des rôles, les WebSockets pour le temps réel, Prometheus, Grafana, Kubernetes, GitHub Actions, les journaux d'audit et un module multi-VM plus avancé.")

    page_break(doc)
    add_title(doc, "BIBLIOGRAPHIE", 1)
    refs = [
        "Documentation officielle React, React Router et Framer Motion.",
        "Documentation officielle FastAPI et Pydantic.",
        "Documentation Docker Engine, Docker Compose et Docker SDK for Python.",
        "Documentation Microsoft Azure Virtual Machines et Azure CLI.",
        "Documentation Nginx, Certbot et Let's Encrypt.",
        "Documentation Terraform CLI et HashiCorp Configuration Language.",
        "Documentation Python psutil pour la collecte de métriques système.",
    ]
    for i, ref in enumerate(refs, 1):
        add_p(doc, f"[{i}] {ref}", style="Normal")

    page_break(doc)
    add_title(doc, "ANNEXES", 1)
    add_title(doc, "Annexe A : Commandes de lancement", 2)
    commands = [
        "Backend local : cd backend && uvicorn main:app --reload --host 0.0.0.0 --port 8000",
        "Frontend local : cd frontend && npm install && npm start",
        "Docker : docker compose up -d --build",
        "Logs backend : docker compose logs -f backend",
        "Arrêt : docker compose down",
    ]
    for cmd in commands:
        add_p(doc, cmd, style="Normal")
    add_title(doc, "Annexe B : Variables d'environnement", 2)
    add_p(doc, "REACT_APP_API_URL=http://localhost:8000 en développement local.")
    add_p(doc, "REACT_APP_API_URL=https://cloudadminyassine.duckdns.org en production.")
    add_title(doc, "Annexe C : Améliorations prévues", 2)
    add_bullets(doc, ["Authentification JWT.", "Gestion des rôles.", "WebSocket temps réel.", "Alertes avancées.", "Prometheus et Grafana.", "Kubernetes.", "CI/CD GitHub Actions.", "Audit logs.", "Gestion multi-VM.", "Module de sauvegarde et restauration."])

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_report()
